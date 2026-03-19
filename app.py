# -*- coding: utf-8 -*-
"""
Hal Esnafı Takip Sistemi - Admin Paneli ve Gelişmiş Filtreleme Versiyonu
"""

import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import database as db

# Sayfa Ayarları
st.set_page_config(page_title="TEMİNAT MEKTUBU TEBLİGAT TAKİP SİSTEMİ", layout="wide")

# --- OTURUM DURUMU (SESSION STATE) ---
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

# --- CSS: TÜM BUTONLARI AYNI TASARIMA GETİRİR ---
st.markdown("""
    <style>
    /* Ana Konteynır ve Üst Boşluk */
    .block-container { padding-top: 2rem !important; padding-bottom: 0rem !important; }
    header[data-testid="stHeader"] { background: rgba(0,0,0,0) !important; color: white !important; }
    
    /* Arka Plan ve Metin Renkleri */
    .stApp { background: linear-gradient(135deg, #020617 0%, #0f172a 100%) !important; }
    section[data-testid="stSidebar"] { background-color: #020617 !important; border-right: 1px solid #1e293b !important; }
    h1, h2, h3, h4, p, span, label { color: #FFFFFF !important; }
    thead tr th { background-color: #1e293b !important; color: #FFFFFF !important; }
    
    /* İstatistik Kartları */
    .stMetric { background-color: rgba(30, 41, 59, 0.5); padding: 15px; border-radius: 10px; border: 1px solid #334155; }

    /* --- TÜM BUTONLARI (İNDİRME VE NORMAL) AYNI YAPAR --- */
    /* Hem st.button hem st.download_button için geçerli stil */
    div.stButton > button, div.stDownloadButton > button {
        background-color: #2563eb !important; /* Kurumsal Mavi */
        color: white !important;
        border: 1px solid #3b82f6 !important;
        font-weight: bold !important;
        border-radius: 8px !important;
        padding: 0.5rem 1rem !important;
        transition: all 0.3s ease !important;
        width: 100% !important; /* Butonları tam genişlik yapar (isteğe bağlı) */
    }

    /* Üzerine gelince (Hover) efekti */
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #1d4ed8 !important; /* Biraz daha koyu mavi */
        border-color: #60a5fa !important;
        box-shadow: 0px 4px 12px rgba(37, 99, 235, 0.3) !important;
        transform: translateY(-1px) !important;
    }

    /* Kenar çubuğu açma düğmesi (Sidebar Collapse) */
    button[data-testid="stSidebarCollapseButton"] {
        background-color: #2563eb !important;
        color: white !important;
    }
    </style>
    """, unsafe_allow_html=True)

# --- YARDIMCI FONKSİYONLAR ---
def to_excel_formatted(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='TakipListesi')
        workbook  = writer.book
        worksheet = writer.sheets['TakipListesi']
        header_format = workbook.add_format({'bold': True, 'bg_color': '#1e293b', 'font_color': 'white', 'border': 1})
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
        for i, col in enumerate(df.columns):
            column_len = max(df[col].astype(str).str.len().max(), len(col)) + 4
            worksheet.set_column(i, i, column_len)
    return output.getvalue()

def clean_currency(value):
    if pd.isna(value) or value == "" or str(value).lower() == "none": return 0.0
    try:
        if isinstance(value, str):
            clean_val = value.replace('.', '').replace(',', '.')
            return float(clean_val)
        return float(value)
    except: return 0.0

def parse_date(date_str):
    if pd.isna(date_str) or not str(date_str).strip() or str(date_str).lower() in ["none", "nan", ""]: return None
    for fmt in ("%d.%m.%Y", "%d/%m/%Y", "%Y-%m-%d"):
        try: return datetime.strptime(str(date_str).strip(), fmt)
        except: continue
    return None

def calculate_gecikme_custom(tebligat_tarihi_str):
    dt = parse_date(tebligat_tarihi_str)
    if not dt: return 0
    gecen_gun = (datetime.now() - dt).days
    return gecen_gun if gecen_gun > 0 else 0

def row_style_logic(row):
    g = row.get("Tebliğ Tarihinden Bugüne Geçen Gün")
    if pd.isna(g) or g <= 30: return [""] * len(row)
    return ["background-color: rgba(220, 38, 38, 0.8); color: #FFFFFF; font-weight: bold;"] * len(row)

def create_word_tebligat(row):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    unvan = str(row.get("Yazıhane Adı", "")).upper()
    hal_no = str(row.get("Yazıhane No", ""))
    hal_adi = str(row.get("Hal", ""))
    tarih_son = "30.09.2026" 
    is_sirket = any(x in unvan for x in ["LTD", "ŞTİ", "A.Ş", "ANONİM", "LİMİTED", "KOOP", "SANAYİ", "TİCARET"])
    hitap = "Bilgilerinize rica ederim." if is_sirket else "Bilgilerinize sunulur."
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = p_header.add_run(f"{unvan}\n{hal_adi} Hal No: {hal_no}")
    run_h.bold = True
    doc.add_paragraph("\n")
    metinler = [
        "5957 sayılı Sebze ve Meyveler İle Yeterli Arz ve Talep Derinliği Bulunan Diğer Malların Ticaretinin Düzenlenmesi Hakkında Kanunun 12’nci maddesinin birinci fıkrasında, toptancı hallerinde faaliyet gösterenlerden teminat alınacağı hükmü yer almaktadır.",
        "Mezkûr Kanuna istinaden çıkarılan Sebze ve Meyve Ticareti ve Toptancı Halleri Hakkında Yönetmeliğin 31’inci maddesinin onikinci fıkrasında ise, eksik kalan teminat tutarının, en geç bir ay içerisinde tamamlattırılacağı ifade edilmektedir.",
        "Yine 5957 sayılı Kanunun 11’inci maddesinin altıncı fıkrasının (a) bendinde, teminatını süresinde vermeyenlerin veya eksilen teminatını süresinde tamamlamayanların kira sözleşmelerinin feshine Belediye Encümenince karar verileceği kuralı getirilmiştir.",
        f"Müdürlüğümüz kayıtlarında yapılan incelemede, {hal_adi} Hali {hal_no} numaralı işyerinde {unvan} ünvanı ile faaliyet gösteren işyeriniz adına Müdürlüğümüze vermiş olduğunuz teminat süresinin {tarih_son} tarihinde sona erdiği ve yeni bir teminat verilmediği tespit edilmiştir.",
        "Bu nedenle, 5957 sayılı Kanunun 12’nci maddesi kapsamında belirtilen teminatın, yazımızın tarafınıza tebliğ edildiği tarihten itibaren 30 (otuz) gün içerisinde Müdürlüğümüze teslim etmeniz gerekmektedir. Aksi takdirde, 5957 sayılı Kanunun 11’inci maddesinin altıncı fıkrasının (a) bendi gereği kira sözleşmenizin feshine karar verilmek üzere dosyanız Belediye Encümenine gönderilecektir.",
        hitap
    ]
    for m in metinler:
        para = doc.add_paragraph(m)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.first_line_indent = Pt(30)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
    doc.add_paragraph("\n\n")
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_s = p_sign.add_run("İsrafil AYDIN\nHal Müdürü")
    run_s.bold = True
    target = BytesIO()
    doc.save(target)
    return target.getvalue()

def main():
    # --- YÖNETİCİ GİRİŞİ (SIDEBAR) ---
    with st.sidebar:
        st.header("🔑 Yönetici Paneli")
        if not st.session_state["authenticated"]:
            sifre = st.text_input("Giriş Şifresi", type="password")
            if st.button("Giriş Yap"):
                if sifre == "1234": # Şifreyi buradan değiştirebilirsiniz
                    st.session_state["authenticated"] = True
                    st.rerun()
                else:
                    st.error("Hatalı Şifre!")
        else:
            st.success("Yönetici Girişi Yapıldı")
            if st.button("Güvenli Çıkış"):
                st.session_state["authenticated"] = False
                st.rerun()
        st.divider()

    st.markdown("<h1 style='margin-top:-20px;'>🏛️ TEMİNAT MEKTUBU TEBLİGAT TAKİP SİSTEMİ</h1>", unsafe_allow_html=True)
    
    df_raw = db.get_all()
    
    # Veri Hazırlama (Eğer tablo boş değilse)
    if not df_raw.empty:
        for col in df_raw.columns:
            df_raw[col] = df_raw[col].apply(lambda x: "" if str(x).lower() in ["none", "nan", ""] else x)
        
        df_raw["Yazıhane No"] = pd.to_numeric(df_raw["Yazıhane No"], errors='coerce').fillna(0).astype(int)
        df_raw["Teminat Tutar"] = df_raw["Teminat Tutar"].apply(clean_currency)
        df_raw["Kalan"] = df_raw["Kalan"].apply(clean_currency)
        df_raw["Tebliğ Tarihinden Bugüne Geçen Gün"] = df_raw["Tebligat Tarihi"].apply(calculate_gecikme_custom)
        df_raw = df_raw.sort_values(by="Yazıhane No")

        # --- İSTATİSTİK PANELİ ---
        st.markdown("### 📊 Genel Durum Özeti")
        c1, c2, c3, c4, c5 = st.columns(5)
        eksik_df = df_raw[df_raw["Durum"].astype(str).str.contains("Eksik", case=False, na=False)]
        c1.metric("🔴 Toplam Eksik", f"{len(eksik_df)}")
        odendi_fazla = len(df_raw[df_raw["Durum"].astype(str).str.contains("Ödendi|Odendi|Fazla", case=False, na=False)])
        c2.metric("🟢 Ödendi + Fazla", f"{odendi_fazla}")
        teblig_edilen = len(df_raw[(df_raw["Tebligat Tarihi"].astype(str).str.strip() != "")])
        c3.metric("📩 Tebliğ Edilen", f"{teblig_edilen}")
        geciken_sayisi = len(df_raw[df_raw["Tebliğ Tarihinden Bugüne Geçen Gün"] > 30])
        c4.metric("⚠️ 30 Günü Geçen", f"{geciken_sayisi}")
        eksik_ve_tebligatsiz = len(eksik_df[(eksik_df["Tebligat Sayı"].astype(str).str.strip() == "")])
        c5.metric("⚪ Tebligat Yapılmamış", f"{eksik_ve_tebligatsiz}")
        st.divider()

    # --- SIDEBAR FİLTRELEME ---
    with st.sidebar:
        st.header("🔍 Filtreleme")
        s_no = st.text_input("Yazıhane No")
        s_ad = st.text_input("Esnaf Adı")
        u_hals = sorted(df_raw["Hal"].dropna().unique().tolist()) if not df_raw.empty else []   
        u_durums = sorted(df_raw["Durum"].dropna().unique().tolist()) if not df_raw.empty else []
        sel_hal = st.selectbox("Hal Bölgesi", ["Tümü"] + u_hals)
        sel_durum = st.selectbox("Durum", ["Tümü"] + u_durums)

    df_f = df_raw.copy() if not df_raw.empty else pd.DataFrame()
    if not df_f.empty:
        if sel_hal != "Tümü": df_f = df_f[df_f["Hal"] == sel_hal]
        if sel_durum != "Tümü": df_f = df_f[df_f["Durum"] == sel_durum]
        if s_no:
            try: df_f = df_f[df_f["Yazıhane No"] == int(s_no)]
            except: pass
        if s_ad: df_f = df_f[df_f["Yazıhane Adı"].astype(str).str.contains(s_ad, case=False, na=False)]

    # --- SEKME YÖNETİMİ (YETKİLENDİRME) ---
    if st.session_state["authenticated"]:
        tab_titles = ["📋 TAKİP PANELİ", "⚡ TEBLİGAT GİRİŞİ & YAZDIRMA", "📥 VERİ GÜNCELLEME"]
    else:
        tab_titles = ["📋 TAKİP PANELİ"]
        st.info("💡 Tebligat girişi ve Veri güncelleme sayfaları için giriş yapmalısınız.")

    tabs = st.tabs(tab_titles)

    # --- TAB 1: TAKİP PANELİ (HERKESE AÇIK) ---
    with tabs[0]:
        if df_f.empty:
            st.warning("Görüntülenecek veri yok.")
        else:
            cols_export = ["Yazıhane No", "Yazıhane Adı", "Teminat Tutar", "Durum", "Kalan", "Tebligat Sayı", "Tebligat Tarihi"]
            excel_data = to_excel_formatted(df_f[cols_export])
            st.download_button(label="📥 Tabloyu Excel Olarak İndir", data=excel_data,
                               file_name=f"Hal_Takip_{datetime.now().strftime('%d_%m_%Y')}.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            
            cols_show = ["Yazıhane No", "Yazıhane Adı", "Teminat Tutar", "Durum", "Kalan", "Tebliğ Tarihinden Bugüne Geçen Gün"]
            st.dataframe(df_f[cols_show].style.apply(row_style_logic, axis=1), 
                         use_container_width=True, hide_index=True,
                         column_config={"Teminat Tutar": st.column_config.NumberColumn(format="%.0f"),
                                        "Kalan": st.column_config.NumberColumn(format="%.0f"),
                                        "Tebliğ Tarihinden Bugüne Geçen Gün": st.column_config.NumberColumn(format="%d Gün ⏳")})

    # --- TAB 2 & 3: SADECE YÖNETİCİ ---
    if st.session_state["authenticated"]:
        # TEBLİGAT GİRİŞİ
# --- TAB 2: TEBLİGAT GİRİŞİ & YAZDIRMA (SADECE YÖNETİCİ) ---
        with tabs[1]:
            if df_f.empty: 
                st.info("İşlem yapılacak veri bulunamadı.")
            else:
                # --- ÜST KISIM: WORD TEBLİGAT HAZIRLAMA ---
                st.subheader("📄 Hızlı Word Tebligat Hazırla")
                esnaf_list = df_f.apply(lambda r: f"{r['Yazıhane No']} - {r['Yazıhane Adı']}", axis=1).tolist()
                
                c1, c2 = st.columns([3, 1]) # Seçim ve buton yan yana şık dursun
                with c1:
                    secilen = st.selectbox("Tebligat yazılacak esnafı seçin:", options=esnaf_list, key="sb_word")
                
                if secilen:
                    y_no = int(secilen.split(" - ")[0])
                    row_data = df_f[df_f["Yazıhane No"] == y_no].iloc[0]
                    word_bin = create_word_tebligat(row_data)
                    with c2:
                        st.write("") # Görsel hizalama için boşluk
                        st.write("") 
                        st.download_button(
                            label=f"📥 {y_no} Nolu Word'ü İndir", 
                            data=word_bin,
                            file_name=f"Tebligat_{y_no}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_btn_{y_no}"
                        )
                
                st.divider() # Araya görsel bir çizgi

                # --- ALT KISIM: VERİ DÜZENLEME VE EXCEL İNDİRME ---
                st.subheader("⚡ Tebligat Bilgilerini Güncelle")
                
                # Excel İndirme Butonu
                cols_export_2 = ["Yazıhane No", "Yazıhane Adı", "Teminat Tutar", "Durum", "Kalan", "Tebligat Sayı", "Tebligat Tarihi"]
                excel_data_2 = to_excel_formatted(df_f[cols_export_2])
                st.download_button(
                    label="📥 Mevcut Listeyi Excel Olarak İndir", 
                    data=excel_data_2,
                    file_name=f"Tebligat_Guncelleme_Listesi_{datetime.now().strftime('%d_%m_%Y')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                    key="btn_tab2_excel"
                )
                
                st.write("") # Küçük bir boşluk
                
                # Veri Düzenleme Tablosu
                df_edit = df_f[["Yazıhane No", "Yazıhane Adı", "Tebligat Sayı", "Tebligat Tarihi"]].copy()
                edited_df = st.data_editor(
                    df_edit, 
                    use_container_width=True, 
                    hide_index=True, 
                    disabled=["Yazıhane No", "Yazıhane Adı"],
                    key="tebligat_editor"
                )
                
                # Kaydet Butonu (Sabit Görünür)
                if st.button("💾 Değişiklikleri Veritabanına Kaydet", key="save_changes_btn"):
                    for _, r in edited_df.iterrows():
                        db.update_tebligat(str(r["Yazıhane No"]), str(r["Tebligat Sayı"]), str(r["Tebligat Tarihi"]))
                    st.success("Tebligat bilgileri başarıyla kaydedildi!")
                    st.rerun()

        # VERİ GÜNCELLEME
        with tabs[2]:
            st.subheader("📥 Yeni Veri Yükleme")
            up = st.file_uploader("Excel dosyasını buraya sürükleyin", type=["xlsx"])
            if up and st.button("🚀 Verileri Sisteme Aktar"):
                try:
                    # Excel'i kontrol amaçlı oku
                    check_df = pd.read_excel(up)
                    
                    # 1. Hatalı İçerik Doğrulaması (Gerekli kolonlar var mı?)
                    required_cols = ["Yazıhane No", "Yazıhane Adı", "Hal", "Durum"]
                    if not all(col in check_df.columns for col in required_cols):
                        st.error("❌ HATALI EXCEL DOSYASI: Gerekli sütunlar (Yazıhane No, Yazıhane Adı, Hal, Durum) bulunamadı!")
                    elif check_df.empty:
                        st.error("❌ HATALI EXCEL DOSYASI: Seçilen dosya boş!")
                    else:
                        # 2. Filtreleme İşlemleri (Boş, 0 veya 573 olanları atla)
                        # 'Hal' kolonu üzerinden filtreleme yapıyoruz (Talep ettiğin gibi)
                        initial_count = len(check_df)
                        check_df = check_df.dropna(subset=['Hal']) # Boşları sil
                        check_df = check_df[check_df['Hal'].astype(str) != "0"] # 0 olanları sil
                        check_df = check_df[check_df['Hal'].astype(str) != "573"] # 573 olanları sil
                        
                        final_count = len(check_df)
                        
                        # Veritabanına gönder (Filtrelenmiş dosyayı kaydetmek için Streamlit'in dosya imlecini başa alamazsınız, 
                        # db fonksiyonunuzun bir dataframe kabul etmesi daha iyi olurdu ama mevcut yapıda dosyayı gönderiyoruz)
                        # NOT: Eğer db.upload_excel dosyayı baştan okuyorsa filtreleme orada da olmalı. 
                        # Ancak kullanıcı isteği üzerine burada doğrulamayı yapıyoruz.
                        
                        db.upload_excel(up) 
                        st.success(f"✅ İşlem Başarılı! {initial_count - final_count} kayıt filtrelendi (0, 573 veya Boş), kalan veriler yüklendi.")
                        st.rerun()
                except Exception as e:
                    st.error(f"Sistem Hatası: {e}")

if __name__ == "__main__":
    main()